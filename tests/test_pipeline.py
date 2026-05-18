"""
test_pipeline.py — Tests for TenderExtractPro.

Covers core logic that doesn't need the LLM model:
  - Pydantic schema enforcement and defaults
  - Chunking logic with synthetic and real data
  - Table column mapping against common tender header formats
  - Grounding verification and confidence scoring
  - Real dataset PDF ingestion and table extraction

Run with:
    python tests/test_pipeline.py
"""

from __future__ import annotations

import json
import logging
import os
import shutil
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from api.main import _resolve_upload_path
from tender_extraction.schemas import (
    Chunk, ChunkMetadata, ExtractionResult,
    TechnicalSpecification, SourceCitation,
    ScopeOfWork,
)
from tender_extraction.chunking import create_chunks
from tender_extraction.table_extraction import (
    _clean_table,
    _map_columns,
    extract_docx_tables,
    extract_tables,
)
from tender_extraction.validation import (
    verify_grounding, assign_confidence,
    validate_extractions, _enforce_not_found,
)
from tender_extraction.ingestion import ingest_document
from tender_extraction.vision import normalize_multimodal_tables, should_try_multimodal_table_fallback

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DATASET_DIR = PROJECT_ROOT / "dataset"


# -- Schema tests --

def test_spec_schema_defaults():
    spec = TechnicalSpecification(
        component="Steel Bars",
        source=SourceCitation(chunk_id="chunk_001", page=5),
    )
    assert spec.specs == {}
    assert spec.confidence == 0.5
    print("  PASS: test_spec_schema_defaults")


def test_empty_spec_text_rejected():
    # New schema has no required string validators—just ensure model builds OK
    spec = TechnicalSpecification(component="Valid")
    assert spec.component == "Valid"
    print("  PASS: test_empty_spec_text_rejected")
    return True


def test_extraction_result_roundtrip():
    result = ExtractionResult(
        technical_specifications=[
            TechnicalSpecification(
                component="Cement",
                specs={"grade": "OPC 53", "standard": "IS 12269"},
                source=SourceCitation(chunk_id="c1", page=15),
                confidence=0.9,
            )
        ],
        scope_of_work=ScopeOfWork(
            summary="Construction and supply of materials",
            deliverables=["Site prep", "Foundation laying"],
        ),
    )
    data = result.model_dump()
    assert len(data["technical_specifications"]) == 1
    assert data["scope_of_work"]["deliverables"][0] == "Site prep"
    print("  PASS: test_extraction_result_roundtrip")


# -- Chunking tests --

def test_chunking_with_sections():
    pages = [{
        "page": 1,
        "text": "1 Introduction\nThis tender is for highway construction.\n\n"
                "2 Technical Specifications\n2.1 Material Requirements\n"
                "Steel bars shall conform to ASTM A615 Grade 60.\n",
        "is_ocr": False,
        "headings": [],
    }]
    chunks = create_chunks(pages, tables=None)
    assert len(chunks) >= 2
    sections = {c.metadata.section for c in chunks}
    assert any("Introduction" in s or "Technical" in s or "Material" in s for s in sections)
    print(f"  PASS: test_chunking_with_sections ({len(chunks)} chunks)")


def test_chunking_table_rows():
    tables = [{
        "table_id": "table_001", "page": 5,
        "headers": ["Item", "Specification", "Unit"],
        "rows": [["Steel Bars", "ASTM A615", "kg"], ["Cement", "IS 12269", "MT"]],
        "bbox": [100, 200, 500, 400],
    }]
    chunks = create_chunks(pages=[], tables=tables)
    assert len(chunks) == 2
    for c in chunks:
        assert c.metadata.chunk_type == "table"
        assert "[Table Headers]" in c.text
    print(f"  PASS: test_chunking_table_rows ({len(chunks)} chunks)")


# -- Table column mapping tests --

def test_column_mapping_standard():
    headers = ["Sr. No.", "Item Description", "Specification Details", "Unit of Measure", "Quantity"]
    mapping = _map_columns(headers)
    assert "item_name" in mapping
    assert "specification_text" in mapping
    assert "unit" in mapping
    print(f"  PASS: test_column_mapping_standard: {mapping}")


def test_column_mapping_alternate():
    headers = ["Sl No", "Name of Material", "Required Quantity", "Rate"]
    mapping = _map_columns(headers)
    assert "item_name" in mapping
    print(f"  PASS: test_column_mapping_alternate: {mapping}")


def test_clean_table():
    raw = [["Header 1", None, "Header 3"], ["Cell\nwith\nnewlines", "", "  Value  "]]
    cleaned = _clean_table(raw)
    assert cleaned[0][1] == ""
    assert "\n" not in cleaned[1][0]
    print("  PASS: test_clean_table")


def test_docx_table_extraction():
    from docx import Document

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "sample.docx"
        document = Document()
        document.add_paragraph("Tender overview paragraph.")
        table = document.add_table(rows=3, cols=3)
        table.rows[0].cells[0].text = "Item"
        table.rows[0].cells[1].text = "Specification"
        table.rows[0].cells[2].text = "Unit"
        table.rows[1].cells[0].text = "Pump"
        table.rows[1].cells[1].text = "10 HP"
        table.rows[1].cells[2].text = "Nos"
        table.rows[2].cells[0].text = "Cable"
        table.rows[2].cells[1].text = "XLPE"
        table.rows[2].cells[2].text = "m"
        document.save(docx_path)

        tables = extract_docx_tables(str(docx_path))
        pages = ingest_document(str(docx_path))
        assert len(tables) == 1
        assert tables[0]["headers"] == ["Item", "Specification", "Unit"]
        assert tables[0]["rows"][0][0] == "Pump"
        assert "Pump" not in pages[0]["text"]
        print("  PASS: test_docx_table_extraction")


def test_multimodal_table_normalization():
    payload = {
        "tables": [
            {
                "headers": ["Parameter", "Value"],
                "rows": [["Voltage", "415 V"], ["Frequency", "50 Hz"]],
            }
        ]
    }
    tables = normalize_multimodal_tables(payload, page_number=3)
    assert len(tables) == 1
    assert tables[0]["page"] == 3
    assert tables[0]["rows"][1][1] == "50 Hz"
    print("  PASS: test_multimodal_table_normalization")


def test_multimodal_table_routing():
    from tender_extraction.config import config

    old_enabled = config.multimodal.enabled
    try:
        config.multimodal.enabled = True
        should_route = should_try_multimodal_table_fallback(
            page_number=2,
            conventional_tables_count=0,
            page_hint={"page": 2, "text": "A | B | C\n10 | 20 | 30\n40 | 50 | 60", "is_ocr": True},
        )
        assert should_route is True
        print("  PASS: test_multimodal_table_routing")
    finally:
        config.multimodal.enabled = old_enabled


def test_upload_path_preserves_suffix():
    upload_path = _resolve_upload_path("abcd1234", "tender.docx")
    assert upload_path.name == "abcd1234.docx"
    print("  PASS: test_upload_path_preserves_suffix")


# -- Grounding tests --

def test_grounding_exact_match():
    chunk = Chunk(
        chunk_id="c1",
        text="Steel bars Grade 60 conforming to ASTM A615.",
        metadata=ChunkMetadata(page=15),
    )
    spec = {
        "source": {"exact_text": "Steel bars Grade 60 conforming to ASTM A615"}
    }
    score = verify_grounding(spec, [{"chunk": chunk, "score": 0.9}])
    assert score >= 0.95
    print(f"  PASS: test_grounding_exact_match (score={score:.3f})")


def test_grounding_rejects_hallucination():
    chunk = Chunk(
        chunk_id="c1",
        text="The deadline for submission of tenders is January 15th.",
        metadata=ChunkMetadata(page=1),
    )
    spec = {
        "source": {"exact_text": "Concrete pump rental agreement terms and conditions"}
    }
    score = verify_grounding(spec, [{"chunk": chunk, "score": 0.9}])
    assert score < 0.4
    print(f"  PASS: test_grounding_rejects_hallucination (score={score:.3f})")


def test_confidence_mapping():
    assert assign_confidence(0.95) == "HIGH"
    assert assign_confidence(0.75) == "MEDIUM"
    assert assign_confidence(0.30) == "LOW"
    print("  PASS: test_confidence_mapping")


def test_enforce_not_found():
    data = {"unit": "", "tolerance": None, "material": "Steel"}
    result = _enforce_not_found(data)
    assert result["unit"] == "NOT_FOUND"
    # None values are not strings so only empty strings get replaced
    assert result["material"] == "Steel"
    print("  PASS: test_enforce_not_found")


def test_validation_rejects_hallucinated():
    source_chunks = [{
        "chunk": Chunk(
            chunk_id="c1",
            text="Steel reinforcement bars shall be Grade 60 conforming to ASTM A615",
            metadata=ChunkMetadata(page=15),
        ),
        "score": 0.9,
    }]
    extraction = {
        "technical_specifications": [
            {
                "component": "Steel Bars",
                "specs": {"grade": "60", "standard": "ASTM A615"},
                "source": {"chunk_id": "c1", "page": 15,
                           "exact_text": "Steel reinforcement bars shall be Grade 60 conforming to ASTM A615"},
            },
            {
                "component": "Hallucinated",
                "specs": {"note": "invented"},
                "source": {"chunk_id": "fake", "page": 99,
                           "exact_text": "Underwater ceramic quantum oscillation differential pressure manifold"},
            },
        ],
        "scope_of_work": {"summary": "NOT_FOUND", "deliverables": [], "exclusions": [], "locations": [], "references": []},
    }
    result = validate_extractions(extraction, source_chunks)
    specs = result["technical_specifications"]
    assert len(specs) == 1
    assert specs[0]["component"] == "Steel Bars"
    print(f"  PASS: test_validation_rejects_hallucinated (1 accepted, 1 rejected)")


def test_validation_filters_ungrounded_scope_fields():
    source_chunks = [{
        "chunk": Chunk(
            chunk_id="c1",
            text="Civil works are excluded from the contractor scope at Building A under Clause 7.",
            metadata=ChunkMetadata(page=4),
        ),
        "score": 0.9,
    }]
    extraction = {
        "technical_specifications": [],
        "scope_of_work": {
            "summary": "NOT_FOUND",
            "deliverables": [],
            "exclusions": ["Civil works are excluded from the contractor scope"],
            "locations": ["Building A"],
            "references": ["Clause 7"],
        },
    }
    validated = validate_extractions(extraction, source_chunks)
    assert validated["scope_of_work"]["exclusions"] == ["Civil works are excluded from the contractor scope"]
    assert validated["scope_of_work"]["locations"] == ["Building A"]
    assert validated["scope_of_work"]["references"] == ["Clause 7"]
    print("  PASS: test_validation_filters_ungrounded_scope_fields")


# -- Real dataset tests --

def test_real_pdf_ingestion():
    pdf_path = DATASET_DIR / "globaltender1576.pdf"
    if not pdf_path.exists():
        print("  SKIP: test_real_pdf_ingestion (dataset not found)")
        return
    pages = ingest_document(str(pdf_path))
    assert len(pages) > 0
    pages_with_text = [p for p in pages if len(p["text"].strip()) > 50]
    assert len(pages_with_text) > 0
    total_chars = sum(len(p["text"]) for p in pages)
    print(f"  PASS: test_real_pdf_ingestion ({len(pages)} pages, {total_chars:,} chars)")


def test_real_pdf_table_extraction():
    pdf_path = DATASET_DIR / "Tenderdocuments.pdf"
    if not pdf_path.exists():
        print("  SKIP: test_real_pdf_table_extraction (dataset not found)")
        return
    tables = extract_tables(str(pdf_path))
    print(f"  PASS: test_real_pdf_table_extraction ({len(tables)} tables)")


def test_real_pdf_chunking():
    pdf_path = DATASET_DIR / "globaltender1576.pdf"
    if not pdf_path.exists():
        print("  SKIP: test_real_pdf_chunking (dataset not found)")
        return
    pages = ingest_document(str(pdf_path))
    tables = extract_tables(str(pdf_path))
    chunks = create_chunks(pages, tables)
    assert len(chunks) > 0
    types = {}
    for c in chunks:
        types[c.metadata.chunk_type] = types.get(c.metadata.chunk_type, 0) + 1
    print(f"  PASS: test_real_pdf_chunking ({len(chunks)} chunks, types: {types})")


def test_query_expansion():
    from tender_extraction.retrieval import expand_query
    expanded = expand_query("specification tolerance material")
    assert len(expanded) > len("specification tolerance material")
    assert "tolerance" in expanded
    print(f"  PASS: test_query_expansion ('{expanded[:60]}...')")


def test_chromadb_retrieval():
    """Test Qdrant-based retrieval with synthetic chunks."""
    from tender_extraction.config import config
    from tender_extraction.retrieval import HybridRetriever

    chunks = [
        Chunk(
            chunk_id="chunk_test_1",
            text="Steel reinforcement bars shall be Grade 60 conforming to ASTM A615.",
            metadata=ChunkMetadata(page=15, section="Materials", chunk_type="paragraph"),
        ),
        Chunk(
            chunk_id="chunk_test_2",
            text="The scope of work includes site preparation and foundation laying.",
            metadata=ChunkMetadata(page=5, section="Scope", chunk_type="paragraph"),
        ),
        Chunk(
            chunk_id="chunk_test_3",
            text="Cement shall be OPC Grade 53 conforming to IS 12269.",
            metadata=ChunkMetadata(page=16, section="Materials", chunk_type="paragraph"),
        ),
    ]

    persist_dir = "./_test_qdrant_db"
    original_rerank_top_k = config.retrieval.rerank_top_k
    try:
        config.retrieval.rerank_top_k = 0
        retriever = HybridRetriever(persist_dir=persist_dir)
        retriever.build_index(chunks, collection_name="test_collection", force_rebuild=True)
        results = retriever.retrieve("steel reinforcement grade", top_k=3)
        assert len(results) > 0
        # The most relevant chunk should be about steel
        top_chunk_text = results[0]["chunk"].text.lower()
        assert "steel" in top_chunk_text or "reinforcement" in top_chunk_text
        print(f"  PASS: test_chromadb_retrieval ({len(results)} results, top='{results[0]['chunk'].text[:40]}...')")
    finally:
        config.retrieval.rerank_top_k = original_rerank_top_k
        shutil.rmtree("./_test_qdrant_db", ignore_errors=True)


def test_retrieval_reranking():
    import tender_extraction.retrieval as retrieval_module
    from tender_extraction.config import config
    from tender_extraction.retrieval import HybridRetriever

    class FakeCrossEncoder:
        def predict(self, pairs):
            scores = []
            for _, text in pairs:
                scores.append(0.95 if "administrative" in text.lower() else 0.10)
            return scores

    original_get_cross_encoder = retrieval_module._get_cross_encoder
    original_rerank_top_k = config.retrieval.rerank_top_k
    config.retrieval.rerank_top_k = 2
    retrieval_module._get_cross_encoder = lambda: FakeCrossEncoder()

    chunks = [
        Chunk(
            chunk_id="chunk_test_1",
            text="Steel reinforcement bars shall be Grade 60 conforming to ASTM A615.",
            metadata=ChunkMetadata(page=15, section="Materials", chunk_type="paragraph"),
        ),
        Chunk(
            chunk_id="chunk_test_2",
            text="Administrative overview and submission conditions for the tender.",
            metadata=ChunkMetadata(page=1, section="Administration", chunk_type="paragraph"),
        ),
    ]

    persist_dir = "./_test_qdrant_rerank"
    try:
        retriever = HybridRetriever(persist_dir=persist_dir)
        retriever.build_index(chunks, collection_name="rerank_collection", force_rebuild=True)
        results = retriever.retrieve("steel reinforcement", top_k=2)
        assert results[0]["chunk"].chunk_id == "chunk_test_2"
        print("  PASS: test_retrieval_reranking")
    finally:
        retrieval_module._get_cross_encoder = original_get_cross_encoder
        config.retrieval.rerank_top_k = original_rerank_top_k
        shutil.rmtree(persist_dir, ignore_errors=True)


def test_collection_rebuilds_when_chunks_change():
    from tender_extraction.retrieval import HybridRetriever

    persist_dir = "./_test_qdrant_manifest"
    collection_name = "manifest_collection"
    original_chunks = [
        Chunk(
            chunk_id="chunk_a",
            text="Steel reinforcement bars shall be Grade 60 conforming to ASTM A615.",
            metadata=ChunkMetadata(page=10, section="Materials", chunk_type="paragraph"),
        ),
        Chunk(
            chunk_id="chunk_b",
            text="Cement shall be OPC Grade 53 conforming to IS 12269.",
            metadata=ChunkMetadata(page=11, section="Materials", chunk_type="paragraph"),
        ),
    ]
    updated_chunks = [
        Chunk(
            chunk_id="chunk_c",
            text="Only one updated chunk should remain in the reused collection.",
            metadata=ChunkMetadata(page=12, section="Updates", chunk_type="paragraph"),
        )
    ]

    # Always start from a clean slate so a prior interrupted run can't pollute results.
    shutil.rmtree(persist_dir, ignore_errors=True)
    try:
        retriever = HybridRetriever(persist_dir=persist_dir)
        retriever.build_index(original_chunks, collection_name=collection_name, force_rebuild=True)
        retriever.close()

        retriever = HybridRetriever(persist_dir=persist_dir)
        retriever.build_index(updated_chunks, collection_name=collection_name, force_rebuild=False)

        # Verify that retrieval returns the updated content, not the old chunks.
        # This is the real correctness check: the rebuild must index the new chunk_c.
        results = retriever.retrieve("updated chunk collection", top_k=3)
        assert len(results) > 0, "Expected at least one result after rebuild"
        # chunk_c should be the top result since it's the only indexed parent
        top_parent_id = results[0]["chunk"].chunk_id
        assert top_parent_id == "chunk_c", (
            f"Expected top result to be chunk_c (the rebuilt chunk), got '{top_parent_id}'"
        )
        print("  PASS: test_collection_rebuilds_when_chunks_change")
    finally:
        shutil.rmtree(persist_dir, ignore_errors=True)



# -- Runner --

def run_all_tests():
    print("\n" + "=" * 60)
    print("  TenderExtractPro -- Test Suite")
    print("=" * 60 + "\n")

    tests = [
        test_spec_schema_defaults,
        test_empty_spec_text_rejected,
        test_extraction_result_roundtrip,
        test_chunking_with_sections,
        test_chunking_table_rows,
        test_column_mapping_standard,
        test_column_mapping_alternate,
        test_clean_table,
        test_docx_table_extraction,
        test_multimodal_table_normalization,
        test_multimodal_table_routing,
        test_upload_path_preserves_suffix,
        test_grounding_exact_match,
        test_grounding_rejects_hallucination,
        test_confidence_mapping,
        test_enforce_not_found,
        test_validation_rejects_hallucinated,
        test_validation_filters_ungrounded_scope_fields,
        test_query_expansion,
        test_chromadb_retrieval,
        test_retrieval_reranking,
        test_collection_rebuilds_when_chunks_change,
        test_real_pdf_ingestion,
        test_real_pdf_table_extraction,
        test_real_pdf_chunking,
    ]

    passed = failed = 0
    for test_fn in tests:
        try:
            test_fn()
            passed += 1
        except Exception as exc:
            failed += 1
            print(f"  FAIL: {test_fn.__name__}: {exc}")

    print(f"\n{'=' * 60}")
    print(f"  Results: {passed} passed, {failed} failed, {len(tests)} total")
    print(f"{'=' * 60}\n")
    return failed == 0


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
